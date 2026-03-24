"""文档解析服务"""

import os
from pathlib import Path
from typing import Dict, Any, Optional


class DocumentParser:
    """文档解析器"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """
        解析文档，提取文本内容和元数据

        Args:
            file_path: 文件路径

        Returns:
            包含文本内容和元数据的字典
        """
        file_type = DocumentParser.get_file_type(file_path)

        if file_type == "pdf":
            return PDFParser.parse(file_path)
        elif file_type == "docx":
            return WordParser.parse(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    @staticmethod
    def get_file_type(file_path: str) -> str:
        """获取文件类型"""
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return "pdf"
        elif ext in [".docx", ".doc"]:
            return "docx"
        else:
            return "unknown"


class PDFParser:
    """PDF文档解析器"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """解析PDF文档"""
        text = ""
        pages = 0
        metadata = {}

        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages = len(pdf_reader.pages)

                # 提取元数据
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                    }

                # 提取文本
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- 第 {page_num + 1} 页 ---\n{page_text}"
                    except Exception as e:
                        print(f"警告: 解析第 {page_num + 1} 页失败: {e}")

        except ImportError:
            try:
                import pypdf

                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    pages = len(pdf_reader.pages)

                    if pdf_reader.metadata:
                        metadata = {
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                        }

                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text += f"\n--- 第 {page_num + 1} 页 ---\n{page_text}"
                        except Exception:
                            pass

            except ImportError:
                raise ImportError("请安装 PDF 解析库: pip install PyPDF2")

        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")

        return {
            "text": text.strip(),
            "pages": pages,
            "metadata": metadata
        }


class WordParser:
    """Word文档解析器"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        text = ""
        metadata = {}

        try:
            from docx import Document

            doc = Document(file_path)

            # 提取元数据
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
            }

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text += paragraph.text + "\n"

            # 提取表格文本
            if doc.tables:
                text += "\n--- 表格内容 ---\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text += row_text + "\n"

        except ImportError:
            raise ImportError("请安装 python-docx 库: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word 文档解析失败: {str(e)}")

        return {
            "text": text.strip(),
            "pages": None,
            "metadata": metadata
        }


def extract_text_from_file(file_path: str) -> str:
    """从文件中提取纯文本内容"""
    result = DocumentParser.parse(file_path)
    return result["text"]


def get_file_info(file_path: str) -> Dict[str, Any]:
    """获取文件信息"""
    file_stat = os.stat(file_path)
    file_type = DocumentParser.get_file_type(file_path)

    return {
        "file_name": os.path.basename(file_path),
        "file_size": file_stat.st_size,
        "file_type": file_type,
    }
